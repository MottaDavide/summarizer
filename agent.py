#!/usr/bin/env python3
"""
Meeting Recap Agent
-------------------
Trascrive un file audio .mp3 con Whisper e genera un recap strutturato
della riunione tramite Gemini 2.5 Flash-Lite (gratuito), verificato su
3 passaggi per massima accuratezza.

Utilizzo:
    python agent.py <percorso_audio.mp3> [--output recap.md] [--model medium]
"""

import sys
import os
import re
import shutil
import argparse
import subprocess
from pathlib import Path
from datetime import datetime

from faster_whisper import WhisperModel
from google import genai
from dotenv import load_dotenv

load_dotenv()

# ---------------------------------------------------------------------------
# Risoluzione PATH per ffmpeg (necessario su Windows con winget)
# ---------------------------------------------------------------------------
def _ensure_ffmpeg_in_path() -> None:
    """
    Whisper chiama ffmpeg tramite subprocess, che cerca l'eseguibile nel PATH
    del processo corrente. Su Windows, winget installa ffmpeg in AppData ma
    non aggiorna il PATH della sessione corrente.
    Questa funzione cerca ffmpeg in posizioni note e lo aggiunge al PATH
    del processo se non è già raggiungibile.
    """
    if shutil.which("ffmpeg"):
        return  # ffmpeg già trovato nel PATH, nulla da fare

    app_root = Path(getattr(sys, "_MEIPASS", Path(__file__).resolve().parent))

    # Posizioni note dove winget/chocolatey/scoop installano ffmpeg su Windows
    candidates = [
        app_root / "ffmpeg" / "bin",
        app_root,
        # winget Gyan.FFmpeg (versione full build)
        Path.home() / "AppData/Local/Microsoft/WinGet/Packages",
        # Fallback: cartella standard manuale
        Path("C:/ffmpeg/bin"),
        Path("C:/Program Files/ffmpeg/bin"),
    ]

    for base in candidates:
        if not base.exists():
            continue
        # Cerca ricorsivamente ffmpeg.exe (max 4 livelli per non andare troppo in profondità)
        if base.is_dir() and (base / "ffmpeg.exe").exists():
            ffmpeg_bin = str(base)
            os.environ["PATH"] = ffmpeg_bin + os.pathsep + os.environ.get("PATH", "")
            print(f"      [ffmpeg] Aggiunto al PATH: {ffmpeg_bin}")
            return
        for ffmpeg_exe in base.rglob("ffmpeg.exe"):
            ffmpeg_bin = str(ffmpeg_exe.parent)
            os.environ["PATH"] = ffmpeg_bin + os.pathsep + os.environ.get("PATH", "")
            print(f"      [ffmpeg] Aggiunto al PATH: {ffmpeg_bin}")
            return

    print("ATTENZIONE: ffmpeg non trovato. Installalo con: winget install Gyan.FFmpeg", file=sys.stderr)

_ensure_ffmpeg_in_path()

# ---------------------------------------------------------------------------
# Configurazione modelli
# ---------------------------------------------------------------------------
WHISPER_MODEL_DEFAULT = "medium"     # Opzioni: tiny, base, small, medium, large
# Gemini 2.5 Flash-Lite: gratuito, 1.000 req/giorno, ottimo per riassunti
# Se il modello non fosse disponibile, prova: "gemini-2.5-flash" o "gemini-2.0-flash"
GEMINI_MODEL = "gemini-2.5-flash-lite"
# Ollama: modello locale usato come fallback quando la quota Gemini è esaurita
# Assicurarsi di aver eseguito: ollama pull mistral
OLLAMA_MODEL = "qwen2.5:14b" #"mistral"
# Formati video supportati per l'estrazione audio
VIDEO_EXTENSIONS = {".mp4", ".mkv", ".avi", ".mov", ".webm", ".m4v"}

# ---------------------------------------------------------------------------
# Prompt
# ---------------------------------------------------------------------------

PROMPT_FIRST_PASS = """Sei un assistente esperto nella sintesi di verbali di riunioni aziendali.

Analizza con attenzione la trascrizione sottostante e crea un recap strutturato.
Usa SOLO le informazioni presenti nella trascrizione: non inventare nulla.
Se qualcosa non è chiaro o incompleto, segnalalo esplicitamente.

Struttura obbligatoria (usa esattamente questi titoli):

## PUNTI CHIAVE DISCUSSI
(elenco puntato dei temi principali affrontati)

## DECISIONI E TAKEAWAY
(decisioni prese, conclusioni importanti, punti da ricordare)

## OPEN POINTS
(questioni aperte, dubbi irrisolti, argomenti da approfondire)

## PROSSIMI PASSI / ACTION ITEMS
(azioni concrete da intraprendere; indica il responsabile se menzionato)

---
TRASCRIZIONE:
{transcript}
---

Recap strutturato:"""

PROMPT_VERIFY_AND_POLISH = """Sei un esperto di verbali aziendali e comunicazione professionale.

Hai davanti una trascrizione originale e una prima bozza di recap.
In un unico passaggio devi:
1. Verificare che ogni punto della bozza sia supportato dalla trascrizione; correggere inesattezze e aggiungere eventuali punti importanti omessi.
2. Formattare il risultato come corpo di email pronta per essere inviata ai partecipanti.
3. Massimo 300 parole. Massimo 4 punti per sezione. Linguaggio semplice e diretto. Evidenzia in **grassetto** le decisioni e i takeaway più importanti.
4. Non riportare la trascrizione.

Formato obbligatorio:

OGGETTO: [oggetto email conciso ed efficace]

---

Buongiorno a tutti,

di seguito il recap della riunione del {meeting_date}.

## Key Points
...

## Decisions and Takeaway
...

## Open Points
...

## Next Steps / Action Items
...

A disposizione per qualsiasi chiarimento.

[Firma placeholder]

---
TRASCRIZIONE ORIGINALE:
{transcript}

---
BOZZA DA VERIFICARE E FORMATTARE:
{recap}
---

Email finale:"""

PROMPT_SIMPLE_RECAP = """Sei un assistente esperto nella sintesi di riunioni aziendali.

Analizza la trascrizione sottostante e produci un breve riassunto in italiano.

Regole:
- Il riassunto deve essere di circa 150-200 parole, in prosa fluente, senza sezioni o titoli.
- Dopo il riassunto, aggiungi esattamente una sezione "Punti chiave:" con al massimo 4 bullet point sintetici.
- NON usare formato email. NON aggiungere oggetto, saluti o firma.
- Usa SOLO le informazioni presenti nella trascrizione.
- Non riportare la trascrizione, solo il riassunto e i punti chiave.

---
TRASCRIZIONE:
{transcript}
---

Riassunto:"""


# ---------------------------------------------------------------------------
# Funzioni
# ---------------------------------------------------------------------------

def get_output_dir(input_path: str) -> Path:
    """Restituisce la cartella di output dedicata al file sorgente."""
    source = Path(input_path)
    output_dir = source.parent / source.stem
    output_dir.mkdir(parents=True, exist_ok=True)
    return output_dir

def extract_audio_from_video(video_path: str) -> str:
    """Estrae la traccia audio da un file video in un .wav persistente (16kHz mono)."""
    video = Path(video_path)
    output_path = get_output_dir(video_path) / f"{video.stem}_extracted.wav"
    cmd = [
        "ffmpeg", "-y",
        "-i", video_path,
        "-q:a", "0",
        "-vn",                  # niente video
        "-acodec", "pcm_s16le", # WAV lossless
        "-ar", "16000",         # 16 kHz (ottimale per Whisper)
        "-ac", "1",             # mono
        str(output_path),
    ]
    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0:
        output_path.unlink(missing_ok=True)
        raise RuntimeError(
            f"ffmpeg: estrazione audio fallita.\n{result.stderr[-800:]}"
        )
    print(f"      Audio estratto salvato in: {output_path}")
    return str(output_path)


def prepare_audio(input_path: str) -> tuple:
    """Prepara il file audio per la trascrizione.

    Se il file è un video, estrae l'audio in un file temporaneo.
    Restituisce (audio_path, is_temp): se is_temp è True, eliminare il file dopo l'uso.
    """
    if Path(input_path).suffix.lower() in VIDEO_EXTENSIONS:
        return extract_audio_from_video(input_path), False
    return input_path, False


def transcribe(audio_path: str, model_name: str, progress_callback=None) -> str:
    """Trascrive il file audio con faster-whisper (CTranslate2, ~4x più veloce su CPU).

    progress_callback(float): chiamata con valore 0.0–1.0 al termine di ogni segmento.
    """
    print(f"[1/4] Caricamento modello Whisper '{model_name}' (faster-whisper, int8)...")
    model = WhisperModel(model_name, device="cpu", compute_type="int8")
    print(f"[1/4] Trascrizione in corso: {audio_path}")
    segments, info = model.transcribe(audio_path, beam_size=5)
    texts = []
    for seg in segments:
        texts.append(seg.text.strip())
        if progress_callback and info.duration > 0:
            progress_callback(min(seg.end / info.duration, 1.0))
    transcript = " ".join(texts).strip()
    source_stem = Path(audio_path).stem.removesuffix("_extracted")
    transcript_path = Path(audio_path).parent / f"{source_stem}_transcript.txt"
    transcript_path.write_text(transcript, encoding="utf-8")
    print(f"      Trascrizione salvata in: {transcript_path}")
    print(f"      Lingua rilevata: {info.language} | Caratteri: {len(transcript)}")
    return transcript


def _is_quota_error(exc: Exception) -> bool:
    """Restituisce True se l'eccezione indica quota esaurita o rate-limit."""
    msg = str(exc).lower()
    return any(k in msg for k in ("429", "quota", "resource_exhausted", "rate limit", "rateerror"))


def call_ollama(prompt: str) -> str:
    """Chiama il modello locale tramite Ollama (fallback quando Gemini è esaurito)."""
    try:
        import ollama
    except ImportError:
        raise RuntimeError(
            "Il pacchetto 'ollama' non è installato.\n"
            "Esegui: pip install ollama\n"
            "e assicurati che Ollama sia in esecuzione: https://ollama.com"
        )
    print(f"      [fallback] Invio a Ollama ({OLLAMA_MODEL})...")
    try:
        response = ollama.generate(model=OLLAMA_MODEL, prompt=prompt)
        return response["response"].strip()
    except Exception as exc:
        raise RuntimeError(
            f"Ollama non raggiungibile o modello '{OLLAMA_MODEL}' non installato.\n"
            f"Esegui: ollama pull {OLLAMA_MODEL}\n"
            f"Dettaglio: {exc}"
        ) from exc


def call_gemini(client: genai.Client, prompt: str) -> str:
    """Invia un prompt a Gemini; in caso di quota esaurita cade su Ollama."""
    #return call_ollama(prompt)
    print(f"      Invio a Gemini ({GEMINI_MODEL})...")
    try:
        response = client.models.generate_content(
            model=GEMINI_MODEL,
            contents=prompt,
        )
        return response.text.strip()
    except Exception as exc:
        if _is_quota_error(exc):
            print(f"      [!] Quota Gemini esaurita — fallback su Ollama ({OLLAMA_MODEL})...")
            return call_ollama(prompt)
        raise 


def generate_recap(transcript: str, meeting_date: str, client: genai.Client) -> tuple:
    """
    Pipeline a 2 passaggi:
      Pass 1 — Generazione iniziale del recap
      Pass 2 — Verifica + formattazione email finale
    """
    # --- Passaggio 1: Prima bozza ---
    print("[2/3] Passaggio 1/2 — Generazione prima bozza...")
    prompt1 = PROMPT_FIRST_PASS.format(transcript=transcript)
    draft1 = call_gemini(client, prompt1)

    # --- Passaggio 2: Verifica + email ---
    print("[3/3] Passaggio 2/2 — Verifica e formattazione email...")
    prompt2 = PROMPT_VERIFY_AND_POLISH.format(
        transcript=transcript,
        recap=draft1,
        meeting_date=meeting_date,
    )
    final = call_gemini(client, prompt2)

    return final, draft1, ""


def generate_simple_recap(transcript: str, client: genai.Client) -> str:
    """Recap semplice: singolo passaggio Gemini (~150-200 parole + max 4 bullet)."""
    prompt = PROMPT_SIMPLE_RECAP.format(transcript=transcript)
    return call_gemini(client, prompt)


def save_output(final_email: str, draft1: str, draft2: str, transcript: str,
                output_path: str, audio_path: str):
    """Salva il recap finale (e opzionalmente le bozze intermedie) su file."""
    p = Path(output_path)

    content = f"""# Meeting Recap
Generato il: {datetime.now().strftime('%Y-%m-%d %H:%M')}
Sorgente audio: {audio_path}

---

{final_email}

---
## APPENDICE: Trascrizione completa

{transcript}
"""
    p.write_text(content, encoding="utf-8")
    print(f"\n✓ Recap salvato in: {p.resolve()}")

    # Salva anche le bozze intermedie in un file separato
    debug_path = p.with_name(p.stem + "_debug.md")
    debug_content = f"""# Debug — Bozze intermedie
Sorgente: {audio_path}

## BOZZA 1 (generazione iniziale)
{draft1}

---

## BOZZA 2 (dopo verifica)
{draft2}
"""
    debug_path.write_text(debug_content, encoding="utf-8")
    print(f"✓ Bozze intermedie salvate in: {debug_path.resolve()}")


def _paragraph_with_bold(doc, text: str, style=None):
    """Aggiunge un paragrafo interpretando **testo** come grassetto Word."""
    kwargs = {"style": style} if style else {}
    p = doc.add_paragraph(**kwargs)
    for i, part in enumerate(re.split(r'\*\*(.+?)\*\*', text)):
        if part:
            p.add_run(part).bold = (i % 2 == 1)
    return p


def _write_full_recap_to_doc(doc, content: str):
    for line in content.splitlines():
        s = line.strip()
        if s.startswith("## "):
            doc.add_heading(s[3:], level=2)
        elif s.startswith("# "):
            doc.add_heading(s[2:], level=1)
        elif s.startswith("OGGETTO:"):
            p = doc.add_paragraph()
            p.add_run(s).bold = True
        elif s.startswith("- ") or s.startswith("* "):
            _paragraph_with_bold(doc, s[2:], style="List Bullet")
        elif s == "---":
            doc.add_paragraph()
        elif s:
            _paragraph_with_bold(doc, s)


def _write_simple_recap_to_doc(doc, content: str):
    in_bullets = False
    for line in content.splitlines():
        s = line.strip()
        if s.lower().startswith("punti chiave"):
            doc.add_heading("Punti chiave", level=2)
            in_bullets = True
        elif in_bullets and (s.startswith("- ") or s.startswith("* ")):
            _paragraph_with_bold(doc, s[2:], style="List Bullet")
        elif s:
            _paragraph_with_bold(doc, s)
            in_bullets = False


def save_docx(
    content: str,
    transcript: str,
    output_path: str,
    audio_path: str,
    mode: str = "full",
    draft1: str = "",
    draft2: str = "",
):
    """Salva il recap in formato .docx (Word)."""
    from docx import Document
    from docx.shared import RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    meta = doc.add_paragraph()
    run = meta.add_run(
        f"Meeting Recap — Generato il: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n"
        f"Sorgente audio: {audio_path}"
    )
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph()

    if mode == "full":
        _write_full_recap_to_doc(doc, content)
    else:
        _write_simple_recap_to_doc(doc, content)

    doc.add_page_break()
    doc.add_heading("Trascrizione completa", level=1)
    doc.add_paragraph(transcript)

    p = Path(output_path)
    doc.save(str(p))
    print(f"\n✓ Recap salvato in: {p.resolve()}")

    if mode == "full" and draft1:
        debug_path = p.with_name(p.stem + "_debug.md")
        debug_content = (
            f"# Debug — Bozze intermedie\nSorgente: {audio_path}\n\n"
            f"## BOZZA 1 (generazione iniziale)\n{draft1}\n\n---\n\n"
            f"## BOZZA 2 (dopo verifica)\n{draft2}\n"
        )
        debug_path.write_text(debug_content, encoding="utf-8")
        print(f"✓ Bozze intermedie salvate in: {debug_path.resolve()}")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description="Trascrivi e riassumi una riunione da file audio o video"
    )
    parser.add_argument("audio", help="Percorso del file audio (.mp3, .wav, ...) o video (.mp4, .mkv, ...)")
    parser.add_argument(
        "--output", "-o",
        default=None,
        help="File di output (default: <nome_audio>_recap.md)",
    )
    parser.add_argument(
        "--model", "-m",
        default=WHISPER_MODEL_DEFAULT,
        choices=["tiny", "base", "small", "medium", "large"],
        help=f"Modello Whisper da usare (default: {WHISPER_MODEL_DEFAULT})",
    )
    parser.add_argument(
        "--docx",
        action="store_true",
        default=False,
        help="Salva l'output come .docx invece di .md",
    )
    args = parser.parse_args()

    audio_path = args.audio
    if not Path(audio_path).exists():
        print(f"Errore: file non trovato: {audio_path}", file=sys.stderr)
        sys.exit(1)

    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        print(
            "Errore: GEMINI_API_KEY non trovata.\n"
            "Crea un file .env con: GEMINI_API_KEY=AIza...\n"
            "Ottieni la chiave gratuita su: https://aistudio.google.com/apikey",
            file=sys.stderr,
        )
        sys.exit(1)

    ext = ".docx" if args.docx else ".md"
    if args.output:
        output_path = args.output
    else:
        output_dir = get_output_dir(audio_path)
        output_path = str(output_dir / f"{Path(audio_path).stem}_recap{ext}")
    meeting_date = datetime.now().strftime("%d/%m/%Y")

    print("=" * 60)
    print("  MEETING RECAP AGENT")
    print("=" * 60)
    print(f"  Audio   : {audio_path}")
    print(f"  Whisper : {args.model}")
    print(f"  Gemini  : {GEMINI_MODEL}  [free tier]")
    print(f"  Output  : {output_path}")
    print("=" * 60)

    # 1. Estrazione audio (se necessario) + trascrizione
    audio_for_whisper, is_temp = prepare_audio(audio_path)
    if is_temp:
        print(f"[0/4] Estrazione audio da video in corso...")
    try:
        transcript = transcribe(audio_for_whisper, args.model)
    finally:
        if is_temp:
            Path(audio_for_whisper).unlink(missing_ok=True)

    # 2. Generazione recap (3 passaggi)
    client = genai.Client(api_key=api_key)
    final_email, draft1, draft2 = generate_recap(transcript, meeting_date, client)

    # 3. Salvataggio
    if args.docx:
        save_docx(final_email, transcript, output_path, audio_path, "full", draft1, draft2)
    else:
        save_output(final_email, draft1, draft2, transcript, output_path, audio_path)

    # 4. Stampa a schermo
    print("\n" + "=" * 60)
    print("  RECAP FINALE")
    print("=" * 60)
    print(final_email)


if __name__ == "__main__":
    main()
